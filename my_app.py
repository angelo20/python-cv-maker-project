from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

#add profile picture to cv document
document.add_picture(
    'profile.jpg', 
    width=Inches(2.0)
)

#get info from user
name = input('Enter Your name? ')
speak('Hello '+ name + 'How are you today')

speak('what is your phone number? ')
phone_number = input('what is your phone number? ')
speak('what is your phone email?  ')
email = input('what is your email? ')

#add text info to document
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# add heading to document
document.add_heading('About Me')
document.add_paragraph(
    input('Tell me about yourself ')
)

# add work experience 
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Enter Company ') 
from_date = input('From date ')
to_date = input('to date ')

p.add_run(company+ ' ').bold = True
p.add_run(from_date +'-' + to_date + '\n').italic = True
experience_details = input('Describe Your experience at '+ company)
p.add_run(experience_details)

# add more experience 
while True:
    has_more = input('Do you have more experiences? Y/N: ')
    if has_more.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company ') 
        from_date = input('From date ')
        to_date = input('to date ')

        p.add_run(company+ ' ').bold = True
        p.add_run(from_date +'-' + to_date + '\n').italic = True
        experience_details = input('Describe Your experience at '+ company+' ')
        p.add_run(experience_details)
    
    else:
        break

# add skills 
document.add_heading('Skills')
skill = input('Enter one of your skills')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more = input('Do you have more skills? Y/N')
    if has_more.lower() == 'yes':
        skill = input('Enter one of your skills')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
  
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Python in collaboration with amigoscode and Institut QuickBooks"
document.save('cv.docx') 